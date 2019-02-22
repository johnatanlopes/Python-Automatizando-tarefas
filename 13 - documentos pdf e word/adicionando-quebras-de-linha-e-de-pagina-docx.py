'''
Quando adicionar a quebra de linha ele finaliza a página que foi adicionado e inicia outra.
'''

import docx
doc = docx.Document()
doc.add_paragraph('This is on the first page!')
doc.paragraphs[0].runs[0].add_break(docx.enum.text.WD_BREAK.PAGE)
doc.add_paragraph('This is on the second page!')
doc.save('twoPage.docx')