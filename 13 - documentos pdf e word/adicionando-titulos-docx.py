'''
Os argumentos de add_heading() são uma string com o texto do titulo e um inteiro de 0 a 4. O deixa com estilo Title, que é usado para o nível mais alto do documento.
Os inteiros de 1 a 4 servem para vários níveis de títulos, em que 1 é o título principal e 4 é o subtitulo de nivel mais baixo. 
'''

import docx

doc = docx.Document()
doc.add_heading('Header 0', 0)
doc.add_heading('Header 1', 1)
doc.add_heading('Header 2', 2)
doc.add_heading('Header 3', 3)
doc.add_heading('Header 4', 4)
doc.save('headings.docx')