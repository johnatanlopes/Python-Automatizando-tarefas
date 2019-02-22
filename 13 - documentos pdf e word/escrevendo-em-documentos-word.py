import docx

# Criando nosso primeiro arquivo
doc = docx.Document()
doc.add_paragraph('Hello world!')
doc.save('helloworld.docx')


# Criando arquivo e adicionando multiplos paragrafos
# Repare que o add_run é como se fosse o append, adicionando um novo trecho ao paragrafo
doc = docx.Document()
doc.add_paragraph('Hello world!')
paraObj1 = doc.add_paragraph('This is a second paragraph')
paraObj2 = doc.add_paragraph('This is a yet another paragraph')
paraObj1.add_run(' This text is being added to the second paragraph')
doc.save('multipleParagraphs.docx')