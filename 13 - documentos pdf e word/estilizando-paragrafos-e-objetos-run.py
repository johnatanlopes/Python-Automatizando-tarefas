'''
O word e outros processadores de texto usam estilos para que a apresentação visual de tipos semelhantes de texto permaneça consistente e fácil de alterar.

Podemos ver no word os estilos pressionando as teclas CTRL-ALT-SHIFT-S

Em documentos Word, há três tipos de estilo: estilos de parágrafo que podem ser aplicados em objetos Paragraph, estilos de caractere podem ser aplicados a objetos 
Run, e estilos vinculados que podem ser aplicados a ambos os tipos de objeto.

Os valores de string para os estilos default do Word são:

normal          Heading5            ListBullet          ListParagraph
BodyText        Heading6            ListBullet2         MacroText
BodyText2       Heading7            ListBullet3         NoSpacing
BodyText3       Heading8            ListContinue        Quote
Caption         Heading9            ListContinue2       Subtitle
Heading1        IntenseQuote        ListContinue3       TOCHeading    
Heading2        List                ListNumber          Title
Heading3        List2               ListNumber2
Heading4        List3               ListNumber3

Ao definir o atributo style não utilize espaços no nome do estilo. Por exemplo, embora o nome do estilo possa ser Subtle Emphasis, defina o atributo style com 
o valor de string 'SubtleEmphasis'.


Atributos de Run

Os runs podem ser estilizados por meio de atributos de text. Cada atributo pode ser definido com um de três valores:
 True (o atributo está sempre habilitado, independentemente de outros estilos aplicados ao run), False (o atributo está sempre desabilitado) ou None 
 (usa qualquer estilo definido no run como default)


Atributos de text do objeto Run
    bold        - O texto aparece em negrito
    italic      - O texto aparece em itálico
    underline   - O texto é sublinhado
    strike      - O texto aparece com uma linha no meio (tachado)
    double_strike       - O texto aparece com duas linhas no meio (tachado duplo)
    all_caps    - O texto aparece em letras maiusculas
    small_caps  - O texto aparece em letras maiusculas e as letras minusculas têm dois pontos a menos.
    shadow      - O texto aparece sombreado
    outline     -    O texto aparece contornado em vez de ser sólido
    rtl         - O texto é escrito da direita para a esquerda
    imprint     - O texto aparece em profundidade na página
    emboss      - O texto aparece em relevo na página
'''

import docx

# Alterando o estilos de demo.docx
doc = docx.Document('demo.docx')
print(doc.paragraphs[0].text)
print(doc.paragraphs[0].style)
doc.paragraphs[0].style = 'Normal'

print(doc.paragraphs[1].text)
print(doc.paragraphs[1].runs[0].text, doc.paragraphs[1].runs[1].text, doc.paragraphs[1].runs[2].text, doc.paragraphs[1].runs[3].text)
doc.paragraphs[1].runs[0].style = 'QuoteChar'
doc.paragraphs[1].runs[1].underline = True
doc.paragraphs[1].runs[3].underline = True
doc.save('restyled.docx')