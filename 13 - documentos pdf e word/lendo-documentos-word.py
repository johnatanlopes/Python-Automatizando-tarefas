# pip3 install python-docx
import docx

doc = docx.Document('demo.docx')

# Quantidade de paragrafos
print(len(doc.paragraphs))

# Exibindo o primeiro paragrafo
print(doc.paragraphs[0].text)

# Exibindo o segundo paragrafo
print(doc.paragraphs[1].text)

'''
Cada objeto paragraph tem um atributo runs que é uma lista de objetos Run. Os objetos Run também têm um atributo text que contém somente o 
texto desse run em particular.
'''
print(doc.paragraphs[1].runs)
print(doc.paragraphs[1].runs[0].text)
print(doc.paragraphs[1].runs[1].text)
print(doc.paragraphs[1].runs[2].text)
print(doc.paragraphs[1].runs[3].text)
print(doc.paragraphs[1].runs[4].text)

print('-----------------')

# Obtendo o arquivo completo
def getText(fileName):
    doc = docx.Document(fileName)
    fullText = []
    for para in doc.paragraphs:
        fullText.append(para.text)
    return '\n'.join(fullText)

print(getText('demo.docx'))